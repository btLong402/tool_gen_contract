from __future__ import annotations

from pathlib import Path

from docx import Document

from contract_automation.services.template_service import render_contract, scan_template_tags


def _create_template_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Contract for {{full_name}}")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Address"
    table.cell(0, 1).text = "{{address}}"
    doc.save(str(path))


def _create_escaped_template_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("Contract for {{full\\_name}}")
    doc.save(str(path))


def test_scan_template_tags_finds_all_tags(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    _create_template_docx(template_path)

    tags = scan_template_tags(template_path)

    assert "full_name" in tags
    assert "address" in tags


def test_render_contract_replaces_placeholders(tmp_path: Path) -> None:
    template_path = tmp_path / "template.docx"
    output_path = tmp_path / "output.docx"
    _create_template_docx(template_path)

    render_contract(
        template_path=template_path,
        data={"full_name": "Nguyen Van A", "address": "Hanoi"},
        output_path=output_path,
    )

    assert output_path.exists()

    rendered = Document(str(output_path))
    full_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    table_text = "\n".join(cell.text for table in rendered.tables for row in table.rows for cell in row.cells)

    assert "Nguyen Van A" in full_text
    assert "Hanoi" in table_text
    assert "{{full_name}}" not in full_text
    assert "{{address}}" not in table_text


def test_scan_and_render_work_with_escaped_underscore_tags(tmp_path: Path) -> None:
    template_path = tmp_path / "escaped_template.docx"
    output_path = tmp_path / "escaped_output.docx"
    _create_escaped_template_docx(template_path)

    tags = scan_template_tags(template_path)
    assert "full_name" in tags

    render_contract(
        template_path=template_path,
        data={"full_name": "Le Thi B"},
        output_path=output_path,
    )

    rendered = Document(str(output_path))
    full_text = "\n".join(paragraph.text for paragraph in rendered.paragraphs)
    assert "Le Thi B" in full_text
    assert "{{full\\_name}}" not in full_text
